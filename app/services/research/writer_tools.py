from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from datetime import datetime
import os
import re


OUTPUT_DIR = "output/research"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── Helpers ──────────────────────────────────────────────────────────────────────

def _sanitize_filename(name: str) -> str:
    """Limpia el nombre del archivo eliminando caracteres inválidos."""
    return re.sub(r'[\\/*?:"<>|]', "", name).strip().replace(" ", "_")


def _parse_markdown_sections(content: str) -> list[dict]:
    """Parsea contenido markdown en secciones con título y texto."""
    sections = []
    current_title = None
    current_body = []

    for line in content.splitlines():
        if line.startswith("## "):
            if current_title:
                sections.append({"title": current_title, "body": "\n".join(current_body).strip()})
            current_title = line[3:].strip()
            current_body = []
        elif line.startswith("# "):
            if current_title:
                sections.append({"title": current_title, "body": "\n".join(current_body).strip()})
            current_title = line[2:].strip()
            current_body = []
        else:
            current_body.append(line)

    if current_title:
        sections.append({"title": current_title, "body": "\n".join(current_body).strip()})

    return sections


# ── Markdown ─────────────────────────────────────────────────────────────────────

@tool
def write_markdown_report(title: str, content: str) -> str:
    """Genera un archivo Markdown (.md) con el reporte de investigación.

    Args:
        title: Título del reporte.
        content: Contenido del reporte en formato Markdown.

    Returns:
        Ruta del archivo generado o descripción del error.
    """
    try:
        filename = _sanitize_filename(title)
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.md")

        header = (
            f"# {title}\n\n"
            f"*Generado por Jarvis — {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n---\n\n"
        )

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(header + content)

        return f"Reporte Markdown generado: {filepath}"
    except Exception as e:
        return f"Error al generar el reporte Markdown: {e}"


# ── DOCX ─────────────────────────────────────────────────────────────────────────

@tool
def write_docx_report(title: str, content: str, references: list[str] | None = None) -> str:
    """Genera un archivo Word (.docx) con el reporte de investigación, incluyendo
    formato profesional y sección de referencias.

    Args:
        title: Título del reporte.
        content: Contenido del reporte en formato Markdown (soporta ## para secciones).
        references: Lista de referencias bibliográficas a incluir al final (opcional).

    Returns:
        Ruta del archivo generado o descripción del error.
    """
    try:
        doc = Document()

        # ── Título ──
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # ── Fecha ──
        date_para = doc.add_paragraph(
            f"Generado por Jarvis — {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph()

        # ── Contenido por secciones ──
        sections = _parse_markdown_sections(content)

        if sections:
            for section in sections:
                doc.add_heading(section["title"], level=1)
                if section["body"]:
                    for paragraph in section["body"].split("\n\n"):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
        else:
            # Sin secciones, agregar como texto plano
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        # ── Referencias ──
        if references:
            doc.add_page_break()
            doc.add_heading("Referencias", level=1)
            for i, ref in enumerate(references, 1):
                ref_para = doc.add_paragraph(style="List Number")
                ref_para.add_run(ref).font.size = Pt(10)

        # ── Guardar ──
        filename = _sanitize_filename(title)
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.docx")
        doc.save(filepath)

        return f"Reporte Word generado: {filepath}"
    except Exception as e:
        return f"Error al generar el reporte Word: {e}"


# ── PDF ──────────────────────────────────────────────────────────────────────────

@tool
def write_pdf_report(title: str, content: str, references: list[str] | None = None) -> str:
    """Genera un archivo PDF con el reporte de investigación con formato profesional.

    Args:
        title: Título del reporte.
        content: Contenido del reporte en formato Markdown (soporta ## para secciones).
        references: Lista de referencias bibliográficas a incluir al final (opcional).

    Returns:
        Ruta del archivo generado o descripción del error.
    """
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # ── Título ──
        pdf.set_font("Helvetica", "B", 20)
        pdf.set_text_color(31, 73, 125)
        pdf.multi_cell(0, 12, title, align="C")
        pdf.ln(4)

        # ── Fecha ──
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(136, 136, 136)
        pdf.cell(0, 6, f"Generado por Jarvis — {datetime.now().strftime('%Y-%m-%d %H:%M')}", align="C")
        pdf.ln(10)

        # ── Línea separadora ──
        pdf.set_draw_color(31, 73, 125)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(8)

        # ── Contenido ──
        sections = _parse_markdown_sections(content)

        if sections:
            for section in sections:
                # Título de sección
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(31, 73, 125)
                pdf.multi_cell(0, 8, section["title"])
                pdf.ln(2)

                # Cuerpo
                if section["body"]:
                    pdf.set_font("Helvetica", "", 11)
                    pdf.set_text_color(30, 30, 30)
                    pdf.multi_cell(0, 6, section["body"])
                    pdf.ln(4)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, content)

        # ── Referencias ──
        if references:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(31, 73, 125)
            pdf.cell(0, 10, "Referencias", ln=True)
            pdf.ln(2)

            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            for i, ref in enumerate(references, 1):
                pdf.multi_cell(0, 6, f"[{i}] {ref}")
                pdf.ln(1)

        # ── Guardar ──
        filename = _sanitize_filename(title)
        filepath = os.path.join(OUTPUT_DIR, f"{filename}.pdf")
        pdf.output(filepath)

        return f"Reporte PDF generado: {filepath}"
    except Exception as e:
        return f"Error al generar el reporte PDF: {e}"


# ── Formato de referencias ────────────────────────────────────────────────────────

@tool
def format_apa_reference(
    authors: str,
    year: str,
    title: str,
    journal: str = "",
    volume: str = "",
    pages: str = "",
    doi: str = "",
    url: str = "",
) -> str:
    """Formatea una referencia bibliográfica en estilo APA 7ma edición.

    Args:
        authors: Autores en formato 'Apellido, N.' separados por coma.
        year: Año de publicación.
        title: Título del artículo o paper.
        journal: Nombre de la revista o journal (opcional).
        volume: Volumen de la publicación (opcional).
        pages: Páginas del artículo, ej: '123-145' (opcional).
        doi: DOI del artículo (opcional).
        url: URL del artículo si no tiene DOI (opcional).

    Returns:
        Referencia formateada en estilo APA.
    """
    try:
        ref = f"{authors} ({year}). {title}."

        if journal:
            ref += f" *{journal}*"
            if volume:
                ref += f", *{volume}*"
            if pages:
                ref += f", {pages}"
            ref += "."

        if doi:
            ref += f" https://doi.org/{doi}"
        elif url:
            ref += f" {url}"

        return ref
    except Exception as e:
        return f"Error al formatear la referencia: {e}"